"""文件读取模块，支持 PDF/DOCX/TXT/MD。"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class LoadedDocument:
    path: Path
    text: str


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    texts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        texts.append(page_text.strip())
    return "\n\n".join(texts)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())


def _read_plain(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def load_documents(paths: Sequence[str | Path]) -> List[LoadedDocument]:
    documents: List[LoadedDocument] = []
    for raw in paths:
        path = Path(raw).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            raise ValueError(f"暂不支持的文件类型: {suffix}")
        if suffix == ".pdf":
            text = _read_pdf(path)
        elif suffix == ".docx":
            text = _read_docx(path)
        else:
            text = _read_plain(path)

        documents.append(LoadedDocument(path=path, text=text))
    return documents
